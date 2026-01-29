from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text, lang_flag=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{lang_flag} {text}" if lang_flag else text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'

def add_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

def add_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Arial'

# ========== РУССКАЯ ВЕРСИЯ ==========
add_title("ДОГОВОР КРАТКОСРОЧНОЙ АРЕНДЫ ЖИЛОГО ПОМЕЩЕНИЯ № {{CONTRACT_NUMBER}}")

add_text("г. {{CITY}}                                                              {{CONTRACT_DATE}}")

add_text("""Мы, нижеподписавшиеся,

АРЕНДОДАТЕЛЬ: {{PARTY_A_NAME}}, ИИН {{PARTY_A_IIN}}, тел: {{PARTY_A_PHONE}},
именуемый в дальнейшем «Арендодатель», с одной стороны, и

АРЕНДАТОР: {{PARTY_B_NAME}}, ИИН {{PARTY_B_IIN}}, удостоверение личности № {{PARTY_B_ID_NUMBER}},
тел: {{PARTY_B_PHONE}}, email: {{PARTY_B_EMAIL}},
именуемый в дальнейшем «Арендатор», с другой стороны,

совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем:""")

add_heading("1. ПРЕДМЕТ ДОГОВОРА")
add_text("""1.1. Арендодатель передаёт, а Арендатор принимает во временное платное пользование жилое помещение по адресу: {{OBJECT_ADDRESS}}

1.2. Жилое помещение передаётся в аренду на срок {{TOTAL_DAYS}} суток:
• Дата и время заселения: {{CHECK_IN_DATE}} с {{CHECK_IN_TIME}} (стандартное время заселения — 14:00)
• Дата и время выселения: {{CHECK_OUT_DATE}} до {{CHECK_OUT_TIME}} (стандартное время выселения — 12:00)

1.3. Арендодатель гарантирует, что является законным владельцем/собственником помещения и имеет право сдавать его в краткосрочную аренду.

1.4. Помимо Арендатора в жилом помещении будут проживать: {{ADDITIONAL_GUESTS}}

1.5. Общее количество проживающих: {{GUESTS_COUNT}} человек.""")

add_heading("2. СТОИМОСТЬ И ПОРЯДОК ОПЛАТЫ")
add_text("""2.1. Стоимость аренды за одни сутки составляет: {{DAILY_RATE}} тенге.

2.2. Полная стоимость проживания за указанные даты: {{TOTAL_AMOUNT}} тенге.

2.3. При бронировании вносится предоплата в размере: {{PREPAYMENT}} тенге.

2.4. Оставшаяся сумма уплачивается Арендатором в полном объёме при заселении в наличном или безналичном порядке.

2.5. При заселении вносится обеспечительный депозит в размере {{DEPOSIT}} тенге за сохранность вверенного имущества на период аренды. Депозит возвращается после выезда и проверки состояния жилого помещения.

2.6. В состав платы за аренду включена компенсация расходов Арендодателя по оплате коммунальных услуг.

2.7. В случае раннего заселения (до {{CHECK_IN_TIME}}) или позднего выселения (после {{CHECK_OUT_TIME}}) сумма платы увеличивается по согласованию сторон.

2.8. Политика отмены бронирования:
• За 3 и более дней до заселения — возврат 100% предоплаты;
• За 1-2 дня до заселения — возврат 50% предоплаты;
• В день заселения — предоплата не возвращается.

2.9. При досрочном выезде по инициативе Арендатора оплаченные денежные средства за оставшиеся дни не возвращаются.""")

add_heading("3. ПРАВА И ОБЯЗАННОСТИ АРЕНДОДАТЕЛЯ")
add_text("""3.1. Арендодатель обязуется:

3.1.1. Предоставить Арендатору в пользование жилое помещение вместе с мебелью, необходимой бытовой техникой, посудой, кухонными принадлежностями и комплектами постельного белья.

3.1.2. Передать Арендатору комплект ключей от жилого помещения и подъезда.

3.1.3. Устранять поломки, аварии и неисправности, произошедшие не по вине Арендатора.

3.1.4. Обеспечить работоспособность всех коммуникаций (вода, электричество, отопление, интернет).

3.2. Арендодатель имеет право:

3.2.1. Осуществлять проверку порядка использования Арендатором жилого помещения и имущества с предварительным уведомлением за 2 часа.

3.2.2. Расторгнуть договор досрочно при грубом нарушении Арендатором условий договора.""")

add_heading("4. ПРАВА И ОБЯЗАННОСТИ АРЕНДАТОРА")
add_text("""4.1. Арендатор обязуется:

4.1.1. Соблюдать правила пользования жилыми помещениями и придомовой территорией, не нарушать права и интересы соседей.

4.1.2. Поддерживать жилое помещение в надлежащем состоянии.

4.1.3. Бережно относиться к имуществу Арендодателя.

4.1.4. Незамедлительно информировать Арендодателя о поломках коммуникаций, авариях и чрезвычайных ситуациях.

4.1.5. При выселении передать Арендодателю полученный комплект ключей.

4.1.6. Сдать жилое помещение в чистом виде, без пятен на поверхностях, без мусора, с чистой посудой.

4.2. В жилом помещении СТРОГО ЗАПРЕЩАЕТСЯ:

4.2.1. Проведение шумных мероприятий и вечеринок. При жалобах соседей — выселение без возврата депозита.

4.2.2. Нарушение тишины с 23:00 до 08:00.

4.2.3. Проживание с домашними животными без письменного согласия Арендодателя. При нарушении — выселение без возврата депозита.

4.2.4. Курение в жилом помещении и на балконе/лоджии. Штраф: {{SMOKING_FINE}} тенге.

4.2.5. Зажигание ароматических свечей и открытого огня.

4.2.6. Перемещение мебели без согласия Арендодателя.

4.2.7. Проживание лиц, не указанных в договоре.

4.2.8. Передача ключей третьим лицам.""")

add_heading("5. ПОРЯДОК ПЕРЕДАЧИ И ВОЗВРАТА ЖИЛОГО ПОМЕЩЕНИЯ")
add_text("""5.1. При заселении Арендатор обязан осмотреть состояние жилого помещения (свет, вода, канализация, газ, окна, работа техники) и сообщить Арендодателю о выявленных дефектах.

5.2. Претензии по состоянию жилого помещения и имущества принимаются в течение 2 часов после заселения посредством WhatsApp, Telegram или электронной почты.

5.3. Техническое и внешнее состояние имущества/мебели/техники на момент передачи считается исправным, если Арендатор не сообщил о дефектах в указанный срок.

5.4. При выселении Арендатор:
• Возвращает все ключи;
• Сдаёт помещение в чистом состоянии;
• Сообщает о любых повреждениях.

5.5. Акт приёма-передачи составляется по требованию любой из сторон.""")

add_heading("6. ОБЕСПЕЧИТЕЛЬНЫЙ ДЕПОЗИТ")
add_text("""6.1. Депозит возвращается в полном объёме в течение 24 часов после выезда при соблюдении всех условий договора.

6.2. Из депозита удерживаются:
• Стоимость повреждённого или утраченного имущества;
• Стоимость дополнительной уборки при сильном загрязнении;
• Штрафы за нарушение условий договора;
• Компенсация за позднее выселение.

6.3. При нарушении правил (шум, животные, курение, грязь) депозит не возвращается.

6.4. При удержании из депозита Арендодатель предоставляет фото/видео доказательства и обоснование суммы.""")

add_heading("7. ОТВЕТСТВЕННОСТЬ СТОРОН")
add_text("""7.1. За неисполнение или ненадлежащее исполнение условий Договора Стороны несут ответственность в соответствии с законодательством Республики Казахстан.

7.2. Арендатор несёт полную материальную ответственность за повреждение жилого помещения или ухудшение его состояния, произошедшие по его вине или вине лиц, проживающих с ним.

7.3. Арендатор отвечает за утерю и порчу имущества Арендодателя.

7.4. При невозможности предоставить помещение по вине Арендодателя, он обязан вернуть 100% оплаты и выплатить компенсацию в размере стоимости одних суток аренды.

7.5. Стороны освобождаются от ответственности при форс-мажоре при условии уведомления другой стороны в течение 24 часов.""")

add_heading("8. РАЗРЕШЕНИЕ СПОРОВ")
add_text("""8.1. Все споры решаются путём переговоров.

8.2. При недостижении согласия сторона направляет письменную претензию. Срок рассмотрения — 7 календарных дней.

8.3. Стороны договорились, что уведомления и иные юридически значимые сообщения могут направляться посредством WhatsApp, Telegram, электронной почты.

8.4. При недостижении согласия в досудебном порядке спор передаётся в суд по месту нахождения жилого помещения.""")

add_heading("9. ПЕРСОНАЛЬНЫЕ ДАННЫЕ")
add_text("""9.1. Арендатор, соглашаясь с условиями настоящего Договора, выражает согласие на обработку и хранение его персональных данных в соответствии с законодательством Республики Казахстан о защите персональных данных.

9.2. К настоящему Договору Арендатор прилагает копию удостоверения личности в качестве подтверждения своей личности.""")

add_heading("10. ПРОСТАЯ ЭЛЕКТРОННАЯ ПОДПИСЬ И ЮРИДИЧЕСКАЯ СИЛА")
add_text("""10.1. Стороны договорились, что настоящий Договор подписывается посредством простой электронной подписи (ПЭП) через сервис 2tick.kz с использованием верификации личности путём направления одноразового кода подтверждения (SMS-код, Telegram-код или голосовой звонок).

10.2. В соответствии со статьёй 152 Гражданского кодекса Республики Казахстан стороны признают, что использование простой электронной подписи (ПЭП) является надлежащим способом подписания настоящего Договора по взаимному соглашению сторон.

10.3. Стороны признают, что Договор, подписанный посредством ПЭП через сервис 2tick.kz, имеет равную юридическую силу с договором, подписанным собственноручно на бумажном носителе.

10.4. Стороны подтверждают, что номера телефонов, указанные в настоящем Договоре, принадлежат им лично и находятся под их контролем. Ввод одноразового кода подтверждения является волеизъявлением стороны и приравнивается к собственноручной подписи.

10.5. Электронный экземпляр Договора, хранящийся в системе 2tick.kz, признаётся сторонами оригиналом Договора и может быть использован в качестве письменного доказательства в суде и государственных органах Республики Казахстан.

10.6. Стороны обязуются не оспаривать действительность настоящего Договора на основании его электронной формы.

10.7. При возникновении спора доказательством является запись в системе 2tick.kz, содержащая идентификатор договора, дату, время и способ верификации.""")

add_heading("11. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ")
add_text("""11.1. Настоящий Договор вступает в силу с момента его подписания и действует до выселения Арендатора из жилого помещения.

11.2. Договор составлен в электронной форме и хранится в системе 2tick.kz.

11.3. Стороны по взаимному согласию могут продлить срок аренды. В этом случае составляется дополнительное соглашение, и Арендатор оплачивает дополнительные сутки проживания.

11.4. Все изменения и дополнения к настоящему Договору действительны при оформлении в письменной или электронной форме и подписании обеими сторонами.

11.5. По вопросам, не урегулированным настоящим Договором, стороны руководствуются законодательством Республики Казахстан.


Договор подписан простой электронной подписью (ПЭП) через сервис 2tick.kz
Уникальный идентификатор договора генерируется системой, дата и время подписания фиксируется системой.""")

add_text("""
ПРИЛОЖЕНИЕ № 1
Стоимость имущества, подлежащего возмещению при повреждении/утрате:

{{INVENTORY_LIST}}
""")

doc.add_page_break()

# ========== КАЗАХСКАЯ ВЕРСИЯ ==========
add_title("ТҰРҒЫН ҮЙ-ЖАЙДЫ ҚЫСҚА МЕРЗІМДІ ЖАЛҒА АЛУ ШАРТЫ № {{CONTRACT_NUMBER}}")

add_text("{{CITY}} қ.                                                              {{CONTRACT_DATE}}")

add_text("""Біз, төменде қол қоюшылар,

ЖАЛҒА БЕРУШІ: {{PARTY_A_NAME}}, ЖСН {{PARTY_A_IIN}}, тел: {{PARTY_A_PHONE}},
бұдан әрі «Жалға беруші» деп аталатын, бір тараптан, және

ЖАЛҒА АЛУШЫ: {{PARTY_B_NAME}}, ЖСН {{PARTY_B_IIN}}, жеке куәлік № {{PARTY_B_ID_NUMBER}},
тел: {{PARTY_B_PHONE}}, email: {{PARTY_B_EMAIL}},
бұдан әрі «Жалға алушы» деп аталатын, екінші тараптан,

бірлесіп «Тараптар» деп аталатын, төмендегілер туралы осы Шартты жасасты:""")

add_heading("1. ШАРТ МӘНІ")
add_text("""1.1. Жалға беруші береді, ал Жалға алушы мына мекенжайдағы тұрғын үй-жайды уақытша ақылы пайдалануға қабылдайды: {{OBJECT_ADDRESS}}

1.2. Тұрғын үй-жай {{TOTAL_DAYS}} тәулікке жалға беріледі:
• Кіру күні мен уақыты: {{CHECK_IN_DATE}} сағат {{CHECK_IN_TIME}} (стандартты кіру уақыты — 14:00)
• Шығу күні мен уақыты: {{CHECK_OUT_DATE}} сағат {{CHECK_OUT_TIME}} дейін (стандартты шығу уақыты — 12:00)

1.3. Жалға беруші үй-жайдың заңды иесі/меншік иесі екенін және оны қысқа мерзімді жалға беруге құқығы бар екенін кепілдендіреді.

1.4. Жалға алушыдан басқа тұрғын үй-жайда тұратындар: {{ADDITIONAL_GUESTS}}

1.5. Тұратын адамдардың жалпы саны: {{GUESTS_COUNT}} адам.""")

add_heading("2. ҚҰНЫ ЖӘНЕ ТӨЛЕМ ТӘРТІБІ")
add_text("""2.1. Бір тәулікке жалдау құны: {{DAILY_RATE}} теңге.

2.2. Көрсетілген күндерге тұрудың толық құны: {{TOTAL_AMOUNT}} теңге.

2.3. Брондау кезінде алдын ала төлем енгізіледі: {{PREPAYMENT}} теңге.

2.4. Қалған сома Жалға алушымен кіру кезінде қолма-қол немесе қолма-қолсыз тәртіппен толық көлемде төленеді.

2.5. Кіру кезінде жалдау кезеңінде сеніп тапсырылған мүліктің сақталуы үшін {{DEPOSIT}} теңге мөлшерінде кепілдік депозит енгізіледі. Депозит шыққаннан кейін және тұрғын үй-жайдың жағдайын тексергеннен кейін қайтарылады.

2.6. Жалдау ақысының құрамына Жалға берушінің коммуналдық қызметтерді төлеу шығындарын өтеу кіреді.

2.7. Ерте кіру ({{CHECK_IN_TIME}} дейін) немесе кеш шығу ({{CHECK_OUT_TIME}} кейін) жағдайында төлем сомасы тараптардың келісімі бойынша ұлғаяды.

2.8. Брондауды болдырмау саясаты:
• Кіруге 3 және одан көп күн қалғанда — алдын ала төлемнің 100% қайтарылады;
• Кіруге 1-2 күн қалғанда — алдын ала төлемнің 50% қайтарылады;
• Кіру күні — алдын ала төлем қайтарылмайды.

2.9. Жалға алушының бастамасы бойынша мерзімінен бұрын шыққан кезде қалған күндер үшін төленген ақша қайтарылмайды.""")

add_heading("3. ЖАЛҒА БЕРУШІНІҢ ҚҰҚЫҚТАРЫ МЕН МІНДЕТТЕРІ")
add_text("""3.1. Жалға беруші міндеттенеді:

3.1.1. Жалға алушыға тұрғын үй-жайды жиһазбен, қажетті тұрмыстық техникамен, ыдыс-аяқпен, ас үй құралдарымен және төсек-орын жиынтықтарымен бірге пайдалануға беру.

3.1.2. Жалға алушыға тұрғын үй-жай мен кіреберістен кілттер жиынтығын беру.

3.1.3. Жалға алушының кінәсінен болмаған сынуларды, апаттарды және ақауларды жою.

3.1.4. Барлық коммуникациялардың (су, электр, жылыту, интернет) жұмыс істеуін қамтамасыз ету.

3.2. Жалға берушінің құқықтары:

3.2.1. 2 сағат бұрын алдын ала хабарлай отырып, Жалға алушының тұрғын үй-жай мен мүлікті пайдалану тәртібін тексеруді жүзеге асыру.

3.2.2. Жалға алушы шарт талаптарын өрескел бұзған жағдайда шартты мерзімінен бұрын бұзу.""")

add_heading("4. ЖАЛҒА АЛУШЫНЫҢ ҚҰҚЫҚТАРЫ МЕН МІНДЕТТЕРІ")
add_text("""4.1. Жалға алушы міндеттенеді:

4.1.1. Тұрғын үй-жайларды және үй маңындағы аумақты пайдалану ережелерін сақтау, көршілердің құқықтары мен мүдделерін бұзбау.

4.1.2. Тұрғын үй-жайды тиісті жағдайда ұстау.

4.1.3. Жалға берушінің мүлігіне ұқыпты қарау.

4.1.4. Коммуникациялардың сынуы, апаттар және төтенше жағдайлар туралы Жалға берушіні дереу хабардар ету.

4.1.5. Шыққан кезде Жалға берушіге алынған кілттер жиынтығын тапсыру.

4.1.6. Тұрғын үй-жайды таза күйінде, беттерде дақсыз, қоқыссыз, ыдыс-аяқ таза күйінде тапсыру.

4.2. Тұрғын үй-жайда ҚАТАҢ ТЫЙЫМ САЛЫНАДЫ:

4.2.1. Шулы іс-шаралар мен кештер өткізу. Көршілердің шағымдары болған жағдайда — депозитсіз шығарылады.

4.2.2. 23:00-ден 08:00-ге дейін тыныштықты бұзу.

4.2.3. Жалға берушінің жазбаша келісімінсіз үй жануарларымен тұру. Бұзған жағдайда — депозитсіз шығарылады.

4.2.4. Тұрғын үй-жайда және балконда/лоджияда темекі шегу. Айыппұл: {{SMOKING_FINE}} теңге.

4.2.5. Хош иісті шырақтар мен ашық от жағу.

4.2.6. Жалға берушінің келісімінсіз жиһаз жылжыту.

4.2.7. Шартта көрсетілмеген адамдардың тұруы.

4.2.8. Кілттерді үшінші тұлғаларға беру.""")

add_heading("5. ТҰРҒЫН ҮЙ-ЖАЙДЫ БЕРУ ЖӘНЕ ҚАЙТАРУ ТӘРТІБІ")
add_text("""5.1. Кіру кезінде Жалға алушы тұрғын үй-жайдың жағдайын (жарық, су, кәріз, газ, терезелер, техниканың жұмысы) тексеруге және анықталған ақаулар туралы Жалға берушіге хабарлауға міндетті.

5.2. Тұрғын үй-жай мен мүліктің жағдайы бойынша шағымдар кіргеннен кейін 2 сағат ішінде WhatsApp, Telegram немесе электрондық пошта арқылы қабылданады.

5.3. Жалға алушы көрсетілген мерзімде ақаулар туралы хабарламаса, беру сәтіндегі мүліктің/жиһаздың/техниканың техникалық және сыртқы жағдайы жарамды деп есептеледі.

5.4. Шыққан кезде Жалға алушы:
• Барлық кілттерді қайтарады;
• Үй-жайды таза күйінде тапсырады;
• Кез келген зақымдар туралы хабарлайды.

5.5. Қабылдау-тапсыру актісі кез келген тараптың талабы бойынша жасалады.""")

add_heading("6. КЕПІЛДІК ДЕПОЗИТ")
add_text("""6.1. Шарттың барлық талаптары орындалған кезде депозит шыққаннан кейін 24 сағат ішінде толық көлемде қайтарылады.

6.2. Депозиттен мыналар ұсталады:
• Зақымдалған немесе жоғалған мүліктің құны;
• Қатты ластану кезінде қосымша тазалау құны;
• Шарт талаптарын бұзғаны үшін айыппұлдар;
• Кеш шығу үшін өтемақы.

6.3. Ережелерді бұзған жағдайда (шу, жануарлар, темекі шегу, кір) депозит қайтарылмайды.

6.4. Депозиттен ұстау кезінде Жалға беруші фото/видео дәлелдер және сома негіздемесін ұсынады.""")

add_heading("7. ТАРАПТАРДЫҢ ЖАУАПКЕРШІЛІГІ")
add_text("""7.1. Шарт талаптарын орындамағаны немесе тиісінше орындамағаны үшін Тараптар Қазақстан Республикасының заңнамасына сәйкес жауапкершілік көтереді.

7.2. Жалға алушы өз кінәсінен немесе онымен бірге тұратын адамдардың кінәсінен болған тұрғын үй-жайдың зақымдануы немесе жағдайының нашарлауы үшін толық материалдық жауапкершілік көтереді.

7.3. Жалға алушы Жалға берушінің мүлігінің жоғалуы мен бүлінуі үшін жауап береді.

7.4. Жалға берушінің кінәсінен үй-жайды бере алмаған жағдайда, ол төлемнің 100% қайтаруға және бір тәулік жалдау құны мөлшерінде өтемақы төлеуге міндетті.

7.5. Форс-мажор жағдайларында 24 сағат ішінде екінші тарапқа хабарлау шартымен тараптар жауапкершіліктен босатылады.""")

add_heading("8. ДАУЛАРДЫ ШЕШУ")
add_text("""8.1. Барлық даулар келіссөздер жолымен шешіледі.

8.2. Келісімге қол жеткізілмеген жағдайда тарап жазбаша талап-арыз жібереді. Қарау мерзімі — 7 күнтізбелік күн.

8.3. Тараптар хабарламалар мен өзге де заңды маңызды хабарламаларды WhatsApp, Telegram, электрондық пошта арқылы жіберуге келісті.

8.4. Сотқа дейінгі тәртіпте келісімге қол жеткізілмеген жағдайда дау тұрғын үй-жайдың орналасқан жері бойынша сотқа беріледі.""")

add_heading("9. ДЕРБЕС ДЕРЕКТЕР")
add_text("""9.1. Жалға алушы осы Шарттың талаптарымен келісе отырып, Қазақстан Республикасының дербес деректерді қорғау туралы заңнамасына сәйкес өзінің дербес деректерін өңдеуге және сақтауға келісімін білдіреді.

9.2. Осы Шартқа Жалға алушы өзінің жеке басын растау үшін жеке куәліктің көшірмесін тіркейді.""")

add_heading("10. ҚАРАПАЙЫМ ЭЛЕКТРОНДЫҚ ҚОЛТАҢБА ЖӘНЕ ЗАҢДЫ КҮШІ")
add_text("""10.1. Тараптар осы Шартқа 2tick.kz сервисі арқылы бір реттік растау кодын (SMS-код, Telegram-код немесе дауыстық қоңырау) жіберу арқылы жеке басын растау арқылы қарапайым электрондық қолтаңба (ҚЭҚ) қолданылып қол қоюға келісті.

10.2. Қазақстан Республикасы Азаматтық кодексінің 152-бабына сәйкес тараптар қарапайым электрондық қолтаңбаны (ҚЭҚ) пайдалану тараптардың өзара келісімі бойынша осы Шартқа қол қоюдың тиісті тәсілі болып табылатынын мойындайды.

10.3. Тараптар 2tick.kz сервисі арқылы ҚЭҚ қолданылып қол қойылған Шарттың қағаз тасығышта өз қолымен қол қойылған шартпен бірдей заңды күші бар екенін мойындайды.

10.4. Тараптар осы Шартта көрсетілген телефон нөмірлері оларға жеке тиесілі және олардың бақылауында екенін растайды. Бір реттік растау кодын енгізу тараптың еркін білдіруі болып табылады және өз қолымен қол қоюға теңестіріледі.

10.5. 2tick.kz жүйесінде сақталған Шарттың электрондық данасы тараптармен Шарттың түпнұсқасы ретінде танылады және Қазақстан Республикасының сотында және мемлекеттік органдарында жазбаша дәлел ретінде пайдаланылуы мүмкін.

10.6. Тараптар осы Шарттың жарамдылығын оның электрондық нысанына негізделіп даулемеуге міндеттенеді.

10.7. Дау туындаған жағдайда дәлелдеме шарт идентификаторын, күнін, уақытын және растау тәсілін қамтитын 2tick.kz жүйесіндегі жазба болып табылады.""")

add_heading("11. ҚОРЫТЫНДЫ ЕРЕЖЕЛЕР")
add_text("""11.1. Осы Шарт қол қойылған сәттен бастап күшіне енеді және Жалға алушы тұрғын үй-жайдан шыққанға дейін қолданылады.

11.2. Шарт электрондық нысанда жасалды және 2tick.kz жүйесінде сақталады.

11.3. Тараптар өзара келісім бойынша жалдау мерзімін ұзарта алады. Бұл жағдайда қосымша келісім жасалады және Жалға алушы қосымша тәуліктерді төлейді.

11.4. Осы Шартқа барлық өзгерістер мен толықтырулар жазбаша немесе электрондық нысанда рәсімделген және екі тарап қол қойған жағдайда жарамды болады.

11.5. Осы Шартта реттелмеген мәселелер бойынша тараптар Қазақстан Республикасының заңнамасын басшылыққа алады.


Шартқа 2tick.kz сервисі арқылы қарапайым электрондық қолтаңба (ҚЭҚ) қолданылып қол қойылды
Шарттың бірегей идентификаторы жүйемен жасалады, қол қою күні мен уақыты жүйемен тіркеледі.""")

add_text("""
№ 1 ҚОСЫМША
Зақымдалған/жоғалған жағдайда өтелуге жататын мүліктің құны:

{{INVENTORY_LIST}}
""")

doc.add_page_break()

# ========== ENGLISH VERSION ==========
add_title("SHORT-TERM RESIDENTIAL RENTAL AGREEMENT № {{CONTRACT_NUMBER}}")

add_text("{{CITY}}                                                              {{CONTRACT_DATE}}")

add_text("""We, the undersigned,

LANDLORD: {{PARTY_A_NAME}}, IIN {{PARTY_A_IIN}}, phone: {{PARTY_A_PHONE}},
hereinafter referred to as "Landlord", on the one hand, and

TENANT: {{PARTY_B_NAME}}, IIN {{PARTY_B_IIN}}, ID number {{PARTY_B_ID_NUMBER}},
phone: {{PARTY_B_PHONE}}, email: {{PARTY_B_EMAIL}},
hereinafter referred to as "Tenant", on the other hand,

collectively referred to as "Parties", have entered into this Agreement as follows:""")

add_heading("1. SUBJECT OF THE AGREEMENT")
add_text("""1.1. The Landlord transfers, and the Tenant accepts for temporary paid use the residential premises located at: {{OBJECT_ADDRESS}}

1.2. The premises are rented for a period of {{TOTAL_DAYS}} days:
• Check-in date and time: {{CHECK_IN_DATE}} from {{CHECK_IN_TIME}} (standard check-in time — 2:00 PM)
• Check-out date and time: {{CHECK_OUT_DATE}} until {{CHECK_OUT_TIME}} (standard check-out time — 12:00 PM)

1.3. The Landlord guarantees that they are the legal owner of the premises and have the right to rent it out on a short-term basis.

1.4. In addition to the Tenant, the following persons will reside in the premises: {{ADDITIONAL_GUESTS}}

1.5. Total number of guests: {{GUESTS_COUNT}} persons.""")

add_heading("2. COST AND PAYMENT PROCEDURE")
add_text("""2.1. Daily rental rate: {{DAILY_RATE}} tenge.

2.2. Total cost of stay for the specified dates: {{TOTAL_AMOUNT}} tenge.

2.3. Advance payment upon booking: {{PREPAYMENT}} tenge.

2.4. The remaining amount is paid by the Tenant in full upon check-in in cash or by bank transfer.

2.5. Upon check-in, a security deposit of {{DEPOSIT}} tenge is paid for the safety of entrusted property during the rental period. The deposit is returned after check-out and inspection of the premises.

2.6. The rental fee includes compensation for the Landlord's utility expenses.

2.7. In case of early check-in (before {{CHECK_IN_TIME}}) or late check-out (after {{CHECK_OUT_TIME}}), the fee increases by agreement of the parties.

2.8. Cancellation policy:
• 3 or more days before check-in — 100% refund of prepayment;
• 1-2 days before check-in — 50% refund of prepayment;
• On the day of check-in — prepayment is non-refundable.

2.9. In case of early departure at the Tenant's initiative, paid funds for remaining days are non-refundable.""")

add_heading("3. LANDLORD'S RIGHTS AND OBLIGATIONS")
add_text("""3.1. The Landlord undertakes to:

3.1.1. Provide the Tenant with the premises together with furniture, necessary household appliances, dishes, kitchen utensils, and bed linen sets.

3.1.2. Hand over to the Tenant a set of keys to the premises and building entrance.

3.1.3. Eliminate breakdowns, accidents, and malfunctions not caused by the Tenant.

3.1.4. Ensure operability of all utilities (water, electricity, heating, internet).

3.2. The Landlord has the right to:

3.2.1. Inspect the Tenant's use of the premises and property with 2 hours prior notice.

3.2.2. Terminate the agreement early in case of gross violation of the agreement terms by the Tenant.""")

add_heading("4. TENANT'S RIGHTS AND OBLIGATIONS")
add_text("""4.1. The Tenant undertakes to:

4.1.1. Comply with the rules of using residential premises and common areas, not violate the rights and interests of neighbors.

4.1.2. Maintain the premises in proper condition.

4.1.3. Take care of the Landlord's property.

4.1.4. Immediately inform the Landlord about breakdowns, accidents, and emergencies.

4.1.5. Return the received set of keys upon check-out.

4.1.6. Return the premises in clean condition, without stains, garbage, and with clean dishes.

4.2. The following is STRICTLY PROHIBITED in the premises:

4.2.1. Hosting noisy events and parties. In case of neighbor complaints — eviction without deposit refund.

4.2.2. Disturbing quiet hours from 11:00 PM to 8:00 AM.

4.2.3. Staying with pets without the Landlord's written consent. If violated — eviction without deposit refund.

4.2.4. Smoking in the premises and on the balcony/loggia. Fine: {{SMOKING_FINE}} tenge.

4.2.5. Lighting aromatic candles and open flames.

4.2.6. Moving furniture without the Landlord's consent.

4.2.7. Accommodation of persons not specified in the agreement.

4.2.8. Handing over keys to third parties.""")

add_heading("5. PROCEDURE FOR TRANSFER AND RETURN OF PREMISES")
add_text("""5.1. Upon check-in, the Tenant must inspect the condition of the premises (lights, water, sewage, gas, windows, appliance operation) and inform the Landlord of any defects found.

5.2. Complaints regarding the condition of the premises and property are accepted within 2 hours after check-in via WhatsApp, Telegram, or email.

5.3. If the Tenant does not report defects within the specified period, the technical and external condition of property/furniture/appliances at the time of transfer is considered proper.

5.4. Upon check-out, the Tenant:
• Returns all keys;
• Leaves the premises in clean condition;
• Reports any damages.

5.5. An acceptance certificate is drawn up at the request of either party.""")

add_heading("6. SECURITY DEPOSIT")
add_text("""6.1. The deposit is returned in full within 24 hours after check-out if all agreement terms are met.

6.2. The following may be deducted from the deposit:
• Cost of damaged or lost property;
• Cost of additional cleaning in case of heavy soiling;
• Fines for violation of agreement terms;
• Compensation for late check-out.

6.3. In case of rule violations (noise, pets, smoking, dirt), the deposit is non-refundable.

6.4. When withholding from the deposit, the Landlord provides photo/video evidence and justification of the amount.""")

add_heading("7. LIABILITY OF THE PARTIES")
add_text("""7.1. For non-performance or improper performance of the Agreement terms, the Parties bear liability in accordance with the legislation of the Republic of Kazakhstan.

7.2. The Tenant bears full financial responsibility for damage to the premises or deterioration of its condition caused by their fault or the fault of persons residing with them.

7.3. The Tenant is responsible for loss and damage to the Landlord's property.

7.4. If the Landlord is unable to provide the premises due to their fault, they are obliged to return 100% of the payment and pay compensation in the amount of one day's rental cost.

7.5. The parties are released from liability in case of force majeure, provided the other party is notified within 24 hours.""")

add_heading("8. DISPUTE RESOLUTION")
add_text("""8.1. All disputes are resolved through negotiations.

8.2. If agreement is not reached, the party sends a written claim. Review period — 7 calendar days.

8.3. The Parties agree that notifications and other legally significant messages may be sent via WhatsApp, Telegram, email.

8.4. If agreement is not reached in the pre-trial procedure, the dispute is referred to court at the location of the premises.""")

add_heading("9. PERSONAL DATA")
add_text("""9.1. By agreeing to the terms of this Agreement, the Tenant consents to the processing and storage of their personal data in accordance with the legislation of the Republic of Kazakhstan on personal data protection.

9.2. The Tenant attaches a copy of their ID to this Agreement as confirmation of their identity.""")

add_heading("10. SIMPLE ELECTRONIC SIGNATURE AND LEGAL FORCE")
add_text("""10.1. The Parties have agreed that this Agreement is signed by means of a simple electronic signature (SES) through the 2tick.kz service using identity verification by sending a one-time confirmation code (SMS code, Telegram code, or voice call).

10.2. In accordance with Article 152 of the Civil Code of the Republic of Kazakhstan, the Parties acknowledge that the use of a simple electronic signature (SES) is a proper method of signing this Agreement by mutual agreement of the Parties.

10.3. The Parties acknowledge that the Agreement signed by SES through the 2tick.kz service has equal legal force with an agreement signed by hand on paper.

10.4. The Parties confirm that the phone numbers specified in this Agreement belong to them personally and are under their control. Entering the one-time confirmation code constitutes the Party's expression of will and is equivalent to a handwritten signature.

10.5. The electronic copy of the Agreement stored in the 2tick.kz system is recognized by the Parties as the original Agreement and may be used as written evidence in court and state authorities of the Republic of Kazakhstan.

10.6. The Parties undertake not to dispute the validity of this Agreement based on its electronic form.

10.7. In case of a dispute, the evidence is the record in the 2tick.kz system containing the agreement identifier, date, time, and verification method.""")

add_heading("11. FINAL PROVISIONS")
add_text("""11.1. This Agreement comes into force from the moment of signing and is valid until the Tenant checks out of the premises.

11.2. The Agreement is made in electronic form and stored in the 2tick.kz system.

11.3. The Parties may extend the rental period by mutual consent. In this case, an additional agreement is drawn up, and the Tenant pays for additional days.

11.4. All changes and additions to this Agreement are valid if made in writing or electronic form and signed by both parties.

11.5. On issues not regulated by this Agreement, the parties shall be guided by the legislation of the Republic of Kazakhstan.


Agreement signed with Simple Electronic Signature (SES) via 2tick.kz service
Unique agreement identifier is generated by the system, date and time of signing is recorded by the system.""")

add_text("""
APPENDIX № 1
Cost of property subject to compensation in case of damage/loss:

{{INVENTORY_LIST}}
""")

# Save the document
doc.save('/app/frontend/public/Договор_посуточной_аренды_НОВЫЙ.docx')
print("Document saved successfully!")
